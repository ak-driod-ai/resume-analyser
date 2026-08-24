import json
import os
import time
from pathlib import Path
from typing import List, Optional

from docx import Document
from dotenv import load_dotenv
from groq import Groq
from pydantic import BaseModel, Field
from pypdf import PdfReader


# ============================================================
# 1. CONFIGURATION
# ============================================================

load_dotenv()

my_api_key = os.getenv("GROQ_API_KEY")

if not my_api_key:
    raise ValueError(
        "API ERROR: GROQ_API_KEY not found in environment variables."
    )

client = Groq(api_key=my_api_key)

# Current model you were already using
model = "openai/gpt-oss-120b"


# ============================================================
# 2. JOB DESCRIPTION
# ============================================================

job_description = """
About American Express:

At American Express, our culture is built on a 175-year history of innovation,
shared values and Leadership Behaviors, and an unwavering commitment to back
our customers, communities, and colleagues.

Business Overview:

Innovation is at the heart of everything we do. Every single day, our
technologists enable our customers around the globe to achieve their goals.

Join the Engineering teams and make an impact right away:

• Develop highly scalable and mission critical applications and platforms
  using modern Cloud Technologies and Programming Languages
  (e.g., Java, Golang, Python, JavaScript)

• Function as core members of agile teams building products and solutions
  on a wide range of technology domains including Distributed Real-time
  Transaction Processing, Big Data and Analytics, AI/ML, omnichannel
  capabilities on native iOS, Android, Social Integration and Services/APIs.

• Work in collaboration with talented engineers in day-to-day activities
  and help in reviewing design, coding, and other Software Development
  Life Cycle tasks.

• Be part of a culture of innovation and experimentation.

Requirements:

• Passion for technology and strong desire to learn
• Curious mind to explore new possibilities
• Strong technical skills
• Problem-solving and analytical abilities
• Proficient programming and coding skills
• Ability to work independently on complex problems
• Strong communication and critical thinking
• Cross-functional teamwork
• Business-first approach
• Creativity and self-sufficiency

Educational Requirements:

• Bachelor's or master's degree in computer science, computer engineering,
  or another related technical discipline.
"""


# ============================================================
# 3. PYDANTIC SCHEMAS
# ============================================================

class JobD(BaseModel):
    role: str = ""
    required_skills: List[str] = Field(default_factory=list)
    preferred_skills: List[str] = Field(default_factory=list)
    all_skills: List[str] = Field(default_factory=list)
    minimum_experience: Optional[float] = None
    education_requirements: List[str] = Field(default_factory=list)
    responsibilities: List[str] = Field(default_factory=list)


class Experience(BaseModel):
    company: Optional[str] = None
    role: Optional[str] = None
    duration: Optional[str] = None
    description: Optional[str] = None
    skills_used: List[str] = Field(default_factory=list)


class Project(BaseModel):
    name: Optional[str] = None
    description: Optional[str] = None
    technologies: List[str] = Field(default_factory=list)


class Resume(BaseModel):
    name: Optional[str] = None
    email: Optional[str] = None
    phone: Optional[str] = None
    total_experience_years: Optional[float] = None

    # Skills explicitly listed OR discovered elsewhere
    skills: List[str] = Field(default_factory=list)

    # Work experience
    experience: List[Experience] = Field(default_factory=list)

    # Projects + technologies used
    projects: List[Project] = Field(default_factory=list)

    education: List[str] = Field(default_factory=list)
    certifications: List[str] = Field(default_factory=list)


class MatchResult(BaseModel):
    score: float
    details: dict


# ============================================================
# 4. PARSE JOB DESCRIPTION
# ============================================================

def parse_job_description(job_description_text: str) -> JobD:

    job_schema = JobD.model_json_schema()

    system_prompt = f"""
You are an expert technical recruiter and job-description parser.

Extract the job description into valid JSON matching this schema:

{json.dumps(job_schema, indent=2)}

IMPORTANT RULES:

1. Extract the job role.

2. Extract ALL technical skills mentioned anywhere in the JD.

3. Do NOT only look for a section called "Skills".

4. Analyze:
   - Responsibilities
   - Requirements
   - Qualifications
   - Technologies
   - Programming languages
   - Frameworks
   - Libraries
   - Databases
   - Cloud technologies
   - APIs
   - AI/ML technologies
   - Tools
   - Platforms

5. Put mandatory/required technologies in required_skills.

6. Put explicitly preferred or nice-to-have technologies in
   preferred_skills.

7. Put the complete unique technical skill set in all_skills.

8. Include technologies mentioned inside responsibilities.

9. Do not invent technologies.

10. Return valid JSON only.
"""

    messages = [
        {
            "role": "system",
            "content": system_prompt
        },
        {
            "role": "user",
            "content": (
                "Analyze this job description:\n\n"
                + job_description_text
            )
        }
    ]

    print("\n" + "=" * 60)
    print("ANALYZING JOB DESCRIPTION")
    print("=" * 60)

    response = client.chat.completions.create(
        model=model,
        messages=messages,
        response_format={"type": "json_object"}
    )

    job_data = json.loads(response.choices[0].message.content)

    job = JobD(**job_data)

    print("\n✓ Job description analyzed")
    print("\nRequired Skills:")
    print(", ".join(job.required_skills))

    print("\nPreferred Skills:")
    print(", ".join(job.preferred_skills))

    print("\nAll JD Skills:")
    print(", ".join(job.all_skills))

    return job


# ============================================================
# 5. PARSE RESUME
# ============================================================

def parse_resume(resume_text: str) -> Resume:

    resume_schema = Resume.model_json_schema()

    system_prompt = f"""
You are an expert technical resume parser.

Extract the candidate information into valid JSON matching this schema:

{json.dumps(resume_schema, indent=2)}

IMPORTANT SKILL EXTRACTION RULES:

1. Extract ALL technical skills explicitly mentioned in the resume.

2. Do NOT only look at the Skills section.

3. Analyze the ENTIRE resume.

4. Look at:

   - Skills section
   - Projects
   - Project descriptions
   - Technologies used in projects
   - Work experience
   - Internships
   - Certifications
   - Tools
   - Platforms
   - Programming languages
   - Frameworks
   - Libraries
   - Databases
   - Cloud technologies
   - Networking technologies

5. For EVERY project:

   - Extract project name.
   - Extract project description.
   - Extract technologies/tools used.
   - Put those technologies in the project's technologies list.

6. If a technology is clearly supported by a project description,
   include it in the candidate's main skills list too.

7. If a technology is clearly supported by work experience,
   include it in the candidate's skills list.

8. Do NOT invent technologies.

9. Normalize obvious variations:

   "C plus plus" -> "C++"
   "PostgreSQL DB" -> "PostgreSQL"
   "Amazon Web Services" -> "AWS"

10. Extract skills even if they appear only inside a project.

11. Return valid JSON only.
"""

    messages = [
        {
            "role": "system",
            "content": system_prompt
        },
        {
            "role": "user",
            "content": (
                "Parse this complete resume:\n\n"
                + resume_text
            )
        }
    ]

    print("\n🧠 Extracting resume information...")

    response = client.chat.completions.create(
        model=model,
        messages=messages,
        response_format={"type": "json_object"}
    )

    raw_output = response.choices[0].message.content

    data = json.loads(raw_output)

    resume = Resume(**data)

    print("✓ Resume parsed")
    print(f"✓ Candidate: {resume.name or 'Unknown'}")

    print("\nExtracted Skills:")
    print(", ".join(resume.skills))

    if resume.projects:
        print("\nProject Technologies:")

        for project in resume.projects:
            print(
                f"  {project.name or 'Unnamed Project'}: "
                f"{', '.join(project.technologies)}"
            )

    return resume


# ============================================================
# 6. STREAMING FINAL MATCH
# ============================================================

def final_score(job: JobD, resume: Resume) -> MatchResult:

    match_schema = MatchResult.model_json_schema()

    prompt = f"""
You are an expert technical recruiter.

Compare the candidate's resume against the job description.

JOB DESCRIPTION:

{job.model_dump_json(indent=2)}


CANDIDATE RESUME:

{resume.model_dump_json(indent=2)}


Perform a comprehensive technical and qualification comparison.

IMPORTANT RULES:

1. Compare the JD against ALL candidate skills.

2. Candidate skills can come from:

   - Skills section
   - Projects
   - Project technologies
   - Work experience
   - Internships
   - Certifications

3. A skill found only inside a project MUST still count as a
   candidate skill.

4. Compare against ALL JD skills, not just required_skills.

5. Give higher importance to required skills.

6. Preferred skills should have lower weight.

7. Identify:

   - Exact skill matches
   - Project-based skill matches
   - Experience-based skill matches
   - Partial/related matches
   - Missing important skills

8. Do not claim a candidate has a technology unless the resume
   provides evidence for it.

9. Semantic equivalents can be considered related matches when
   appropriate.

10. Consider:

    - Technical skills
    - Project experience
    - Work experience
    - Education
    - Certifications
    - Responsibilities

11. Calculate an overall score from 0 to 100.

12. Provide a realistic final hiring-oriented verdict.

Return valid JSON matching this schema:

{json.dumps(match_schema, indent=2)}

The "details" object MUST contain:

- candidate_name
- matching_skills
- project_matching_skills
- experience_matching_skills
- missing_important_skills
- partial_matches
- skill_match_percentage
- final_verdict
- explanation
"""

    print("\n" + "=" * 60)
    print("🔍 MATCHING RESUME WITH JOB DESCRIPTION")
    print("=" * 60)

    print("\n🤖 AI analysis:\n")

    # STREAMING ENABLED
    stream = client.chat.completions.create(
        model=model,
        messages=[
            {
                "role": "user",
                "content": prompt
            }
        ],
        response_format={"type": "json_object"},
        stream=True
    )

    full_response = ""

    for chunk in stream:

        if not chunk.choices:
            continue

        content = chunk.choices[0].delta.content

        if content:

            print(content, end="", flush=True)

            full_response += content

    print("\n")

    # Convert streamed JSON into Python object
    data = json.loads(full_response)

    return MatchResult(**data)


# ============================================================
# 7. READ PDF
# ============================================================

def read_pdf(file_path):

    reader = PdfReader(file_path)

    text = ""

    for page in reader.pages:

        page_text = page.extract_text()

        if page_text:
            text += page_text + "\n"

    return text


# ============================================================
# 8. READ DOCX
# ============================================================

def read_docx(file_path):

    document = Document(file_path)

    text = ""

    # Normal paragraphs
    for paragraph in document.paragraphs:

        if paragraph.text.strip():

            text += paragraph.text + "\n"

    # Tables
    for table in document.tables:

        for row in table.rows:

            for cell in row.cells:

                if cell.text.strip():

                    text += cell.text + "\n"

    return text


# ============================================================
# 9. READ RESUME FILE
# ============================================================

def read_resume(file_path):

    file_path = Path(file_path)

    extension = file_path.suffix.lower()

    if extension == ".pdf":

        return read_pdf(file_path)

    elif extension == ".docx":

        return read_docx(file_path)

    else:

        return None


# ============================================================
# 10. MAIN PROGRAM
# ============================================================

def main():

    # --------------------------------------------------------
    # Parse JD once
    # --------------------------------------------------------

    job = parse_job_description(job_description)

    # --------------------------------------------------------
    # Resume folder
    # --------------------------------------------------------

    resume_folder = Path("resumes")

    if not resume_folder.exists():

        print("\nERROR: 'resumes' folder not found.")

        return

    if not resume_folder.is_dir():

        print("\nERROR: 'resumes' is not a directory.")

        return

    all_results = []

    # --------------------------------------------------------
    # Process every resume
    # --------------------------------------------------------

    resume_files = [
        file_path
        for file_path in resume_folder.iterdir()
        if file_path.suffix.lower() in [".pdf", ".docx"]
    ]

    if not resume_files:

        print("\nNo PDF or DOCX resumes found inside 'resumes'.")

        return

    print("\n" + "=" * 60)
    print(f"FOUND {len(resume_files)} RESUME(S)")
    print("=" * 60)

    for file_path in resume_files:

        print("\n")
        print("=" * 60)
        print(f"PROCESSING: {file_path.name}")
        print("=" * 60)

        # ----------------------------------------------------
        # Read resume
        # ----------------------------------------------------

        resume_text = read_resume(file_path)

        if not resume_text:

            print("Could not extract text from resume.")

            continue

        # ----------------------------------------------------
        # Parse resume
        # ----------------------------------------------------

        parsed_resume = parse_resume(resume_text)

        time.sleep(1)

        # ----------------------------------------------------
        # Match resume against JD
        # ----------------------------------------------------

        result = final_score(job, parsed_resume)

        time.sleep(1)

        candidate_name = (
            parsed_resume.name
            or file_path.stem
        )

        all_results.append(
            {
                "filename": file_path.name,
                "name": candidate_name,
                "score": result.score,
                "details": result.details
            }
        )

    # ========================================================
    # SORT RESULTS
    # ========================================================

    all_results.sort(
        key=lambda candidate: candidate["score"],
        reverse=True
    )

    # ========================================================
    # FINAL RESULTS
    # ========================================================

    print("\n")
    print("=" * 70)
    print("                 ALL EVALUATED RESUMES")
    print("=" * 70)

    for index, candidate in enumerate(all_results, 1):

        print(
            f"\n[{index}] "
            f"{candidate['name']} "
            f"({candidate['filename']})"
        )

        print(
            f"Match Score: "
            f"{candidate['score']}/100"
        )

        details = candidate["details"]

        print(
            "\nMatching Skills:"
        )

        for skill in details.get(
            "matching_skills",
            []
        ):

            print(f"  ✓ {skill}")

        print(
            "\nProject Matching Skills:"
        )

        for skill in details.get(
            "project_matching_skills",
            []
        ):

            print(f"  ✓ {skill}")

        print(
            "\nExperience Matching Skills:"
        )

        for skill in details.get(
            "experience_matching_skills",
            []
        ):

            print(f"  ✓ {skill}")

        print(
            "\nPartial Matches:"
        )

        for skill in details.get(
            "partial_matches",
            []
        ):

            print(f"  ~ {skill}")

        print(
            "\nMissing Important Skills:"
        )

        for skill in details.get(
            "missing_important_skills",
            []
        ):

            print(f"  ✗ {skill}")

        print(
            "\nFinal Verdict:"
        )

        print(
            details.get(
                "final_verdict",
                "Not available"
            )
        )

        print("\nExplanation:")

        print(
            details.get(
                "explanation",
                "Not available"
            )

        )

        print("-" * 70)

    # ========================================================
    # TOP CANDIDATES
    # ========================================================

    if all_results:

        print("\n" + "=" * 70)
        print("                    RANKING")
        print("=" * 70)

        for index, candidate in enumerate(
            all_results,
            1
        ):

            print(
                f"{index}. "
                f"{candidate['name']} "
                f"→ "
                f"{candidate['score']}/100"
            )


# ============================================================
# 11. PROGRAM ENTRY POINT
# ============================================================

if __name__ == "__main__":

    main()